from docx import Document
from io import BytesIO


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract clean text from a DOCX file.
    Supports paragraphs + table cells.
    """

    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        print(f"[DOCX ERROR] Failed to read DOCX: {e}")
        return ""

    text_parts = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract text from tables 
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    # Combine and clean text
    if not text_parts:
        return ""

    full_text = " ".join(text_parts)
    clean_text = " ".join(full_text.split())  # normalize spaces

    return clean_text
